#!/usr/bin/env python3
"""
Create a Word document template for lab members to fill out their information
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_field(doc, label, description, example="", required=True):
    """Add a field to fill out"""
    p = doc.add_paragraph()

    # Label (bold)
    run = p.add_run(label)
    run.bold = True

    # Required marker
    if required:
        req = p.add_run(" (REQUIRED)")
        req.font.color.rgb = RGBColor(255, 0, 0)
        req.font.size = Pt(9)
    else:
        opt = p.add_run(" (optional)")
        opt.font.color.rgb = RGBColor(128, 128, 128)
        opt.font.size = Pt(9)

    # Description
    p = doc.add_paragraph(description, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

    # Example
    if example:
        p_ex = doc.add_paragraph(f"Example: {example}", style='List Bullet')
        p_ex.paragraph_format.left_indent = Inches(0.5)
        run = p_ex.runs[0]
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Fill-in line
    p_fill = doc.add_paragraph("Your answer: _________________________________________________")
    p_fill.paragraph_format.left_indent = Inches(0.5)
    p_fill.paragraph_format.space_after = Pt(12)

    return p_fill

def create_template():
    doc = Document()

    # Title
    title = doc.add_heading('CAMER Lab Members Page - Information Form', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_paragraph(
        "Please fill out this form with your information for the CAMER Lab website's people page. "
        "Your information will be displayed alongside your publications and project affiliations."
    )

    doc.add_paragraph(
        "Fields marked as REQUIRED must be filled out. Optional fields can be left blank if not applicable."
    )

    doc.add_paragraph().add_run("").bold = True

    # Section 1: Basic Information
    add_heading(doc, "1. Basic Information", level=1)

    add_field(
        doc,
        "Full Name",
        "Your full name as you want it to appear on the website",
        "John Smith",
        required=True
    )

    add_field(
        doc,
        "Email Address",
        "Your UW-Madison email address (or current institution email for alumni)",
        "jsmith@wisc.edu",
        required=True
    )

    add_field(
        doc,
        "Status",
        "Select one: Graduate Student, Undergraduate Student, or Alumni",
        "Graduate Student",
        required=True
    )

    # For alumni only
    p_alumni = doc.add_paragraph()
    run = p_alumni.add_run("Current Position")
    run.bold = True
    run_opt = p_alumni.add_run(" (REQUIRED for alumni only)")
    run_opt.font.color.rgb = RGBColor(255, 0, 0)
    run_opt.font.size = Pt(9)

    p = doc.add_paragraph(
        "If you are an alumnus/alumna, provide your current position and institution",
        style='List Bullet'
    )
    p.paragraph_format.left_indent = Inches(0.5)

    p_ex = doc.add_paragraph(
        "Example: Assistant Professor, University of Wisconsin-Madison",
        style='List Bullet'
    )
    p_ex.paragraph_format.left_indent = Inches(0.5)
    run = p_ex.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    p_fill = doc.add_paragraph("Your answer: _________________________________________________")
    p_fill.paragraph_format.left_indent = Inches(0.5)
    p_fill.paragraph_format.space_after = Pt(12)

    # Section 2: Research Information (Students Only)
    doc.add_page_break()
    add_heading(doc, "2. Research Information (Graduate/Undergraduate Students Only)", level=1)

    p_note = doc.add_paragraph(
        "NOTE: If you are an alumnus/alumna, skip this section."
    )
    p_note.runs[0].italic = True
    p_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    add_field(
        doc,
        "Research Interest",
        "A brief statement (1-2 sentences) describing your research interests",
        "Investigating how moral appeals in health communication shape public attitudes and behavior change across digital platforms.",
        required=True
    )

    p_proj = doc.add_paragraph()
    run = p_proj.add_run("Associated Projects")
    run.bold = True
    run_opt = p_proj.add_run(" (optional)")
    run_opt.font.color.rgb = RGBColor(128, 128, 128)
    run_opt.font.size = Pt(9)

    p = doc.add_paragraph(
        "List the project IDs you are involved with (e.g., 1_project, 2_project, 3_project). "
        "These correspond to project pages on the website. Leave blank if unsure.",
        style='List Bullet'
    )
    p.paragraph_format.left_indent = Inches(0.5)

    p_ex = doc.add_paragraph(
        "Example: 1_project, 3_project",
        style='List Bullet'
    )
    p_ex.paragraph_format.left_indent = Inches(0.5)
    run = p_ex.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    p_fill = doc.add_paragraph("Your answer: _________________________________________________")
    p_fill.paragraph_format.left_indent = Inches(0.5)
    p_fill.paragraph_format.space_after = Pt(12)

    # Section 3: Photo
    doc.add_page_break()
    add_heading(doc, "3. Profile Photo", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Profile Photo")
    run.bold = True
    run_opt = p.add_run(" (REQUIRED)")
    run_opt.font.color.rgb = RGBColor(255, 0, 0)
    run_opt.font.size = Pt(9)

    doc.add_paragraph(
        "Please provide a professional headshot photo (ideally square, at least 300x300 pixels). "
        "Email your photo to the lab website administrator with the filename: "
        "firstname-lastname.jpg (e.g., john-smith.jpg)"
    )

    p_check = doc.add_paragraph("☐ I have sent my photo to the website administrator")
    p_check.paragraph_format.left_indent = Inches(0.5)

    # Section 4: Online Profiles & Links
    doc.add_page_break()
    add_heading(doc, "4. Online Profiles & Links (All Optional)", level=1)

    doc.add_paragraph(
        "Provide links to your professional profiles. Leave blank if you don't have an account "
        "or prefer not to include it."
    )

    links_info = [
        ("Personal Website", "Your personal or professional website URL", "https://www.yourname.com"),
        ("CV/Resume (PDF)", "Path to your CV PDF file on the website, or full URL", "/assets/pdf/members/john-smith-cv.pdf"),
        ("Google Scholar", "Full URL to your Google Scholar profile", "https://scholar.google.com/citations?user=abc123"),
        ("GitHub", "Full URL to your GitHub profile", "https://github.com/yourusername"),
        ("OSF (Open Science Framework)", "Full URL to your OSF profile", "https://osf.io/abc12/"),
        ("LinkedIn", "Full URL to your LinkedIn profile", "https://linkedin.com/in/yourprofile"),
        ("Twitter/X", "Full URL to your Twitter/X profile", "https://twitter.com/yourusername"),
    ]

    for label, desc, example in links_info:
        add_field(doc, label, desc, example, required=False)

    # Section 5: Publications
    doc.add_page_break()
    add_heading(doc, "5. Publications", level=1)

    doc.add_paragraph(
        "Your publications will be AUTOMATICALLY populated from the lab's bibliography file (papers.bib). "
        "Publications where you are a co-author will automatically appear on your profile."
    )

    doc.add_paragraph(
        "Action Required:"
    ).runs[0].bold = True

    doc.add_paragraph(
        "1. Verify that all your CAMER publications are in the papers.bib file",
        style='List Number'
    )

    doc.add_paragraph(
        "2. Ensure your name in papers.bib matches your name on this form exactly",
        style='List Number'
    )

    doc.add_paragraph(
        "3. If any publications are missing, notify the website administrator",
        style='List Number'
    )

    p_note = doc.add_paragraph(
        "Note: The system handles various name formats (e.g., 'John Smith' vs 'Smith, John' vs 'Smith, J.')"
    )
    p_note.runs[0].italic = True
    p_note.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Section 6: Submission
    doc.add_page_break()
    add_heading(doc, "6. Submission", level=1)

    doc.add_paragraph(
        "When you have completed this form:"
    )

    doc.add_paragraph(
        "1. Save this document with your name: Lastname_Firstname_MemberInfo.docx",
        style='List Number'
    )

    doc.add_paragraph(
        "2. Email the completed form and your profile photo to: [ADMIN EMAIL HERE]",
        style='List Number'
    )

    doc.add_paragraph(
        "3. Your information will be added to the website within 1-2 business days",
        style='List Number'
    )

    # Footer
    doc.add_paragraph()
    p_footer = doc.add_paragraph(
        "Thank you for contributing to the CAMER Lab website!"
    )
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.runs[0].italic = True
    p_footer.runs[0].font.size = Pt(10)

    # Save
    output_path = "/Users/sijiayang/Documents/sijiayangcamer.github.io/CAMER_Member_Information_Form.docx"
    doc.save(output_path)
    print(f"✓ Template created: {output_path}")

if __name__ == '__main__':
    create_template()
